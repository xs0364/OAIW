import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# Deep scan for ALL circled numbers in the filled DOCX
import xml.etree.ElementTree as ET
from docx import Document

doc = Document(r"D:\OAIW\dianfangbaohan\2\result_v2.docx")

# All possible circled number ranges
import re

print("=== Raw XML scan for circled/label numbers ===")
for i, p in enumerate(doc.paragraphs):
    t = p.text
    # Check for ANY label-like characters
    labels = []
    for ch in t:
        cp = ord(ch)
        # CJK Circled numbers: ①=9312 to ⑳=9331, ⑴=9332 to ⒇=9351
        if 9312 <= cp <= 9351:
            labels.append(ch)
        # Regular circled: ⒜=9398 etc
        elif 9398 <= cp <= 9423:
            labels.append(ch)
    if labels:
        print(f"  [{i:2d}] labels={''.join(labels)} in: {t[:150]}")

# ALSO check runs individually (some labels might be in formatting runs)
print("\n=== Per-run text scan for labels ===")
for i, p in enumerate(doc.paragraphs):
    for j, r in enumerate(p.runs):
        for ch in r.text:
            cp = ord(ch)
            if 9312 <= cp <= 9351:
                print(f"  [{i}].run[{j}] has label {ch}: {r.text[:100]}")
                break

# Also check raw XML
print("\n=== XML fragment scan ===")
ns = {'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}
for i, p in enumerate(doc.paragraphs):
    xml_str = ET.tostring(p._element, encoding='unicode')
    # Look for any circled number hex codes
    # ①=0x2460 to ⑳=0x2463, ⑴=0x2474 to ⒇=0x2487
    for code_hex in ['2460','2461','2462','2463','2464','2465','2466','2467',
                     '2474','2475','2476','2477','2478','2479','247A','247B']:
        if code_hex in xml_str.upper():
            print(f"  [{i}] hex {code_hex} found in XML")
            break
