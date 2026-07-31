"""
OAIW RPA 任务路由
"""
from __future__ import annotations

import asyncio
import json
import os
import queue as _queue
import re
import shutil
import uuid
from typing import Optional

import httpx
from fastapi import APIRouter, Depends, Header, HTTPException, UploadFile, File, Form
from fastapi.responses import StreamingResponse
from pydantic import BaseModel
from sqlalchemy.orm import Session

from backend.config import settings
from backend.database import get_db
from backend.core.services import get_current_user
from backend.parser import extract_text

router = APIRouter(prefix="/api/rpa", tags=["rpa"])

# In-memory store for filled DOCX files (download_id -> path)
_telex_download_store: dict[str, str] = {}


class RpaTaskRequest(BaseModel):
    task_type: str  # port_query | port_status | track_cargo | generate_letter | merge_docs
    params: dict = {}


class RpaTaskResponse(BaseModel):
    success: bool
    data: str = ""
    error: str = ""


def _send_completion_notification(user, task_name: str, result: dict):
    """异步发送 RPA 任务完成通知邮件。"""
    try:
        config = json.loads(user.notify_config or "{}")
        if not config.get("on_rpa"):
            return
        to_email = config.get("email", user.email)
        if not to_email:
            return

        status_icon = "&#9989;" if result.get("success") else "&#10060;"
        status_text = "完成" if result.get("success") else "失败"
        data_preview = result.get("data", "")

        from backend.utils.email import send_notification_async
        send_notification_async(
            to_email=to_email,
            subject=f"RPA 任务 {status_text}: {task_name}",
            content_text=(
                f"RPA 任务: {task_name}\n"
                f"状态: {status_text}\n"
                f"{'─' * 40}\n"
                f"{data_preview}"
            ),
            task_name=task_name,
        )
    except Exception:
        pass  # 通知失败不影响主流程


@router.post("/sms/submit")
async def submit_sms_code(data: dict):
    """
    用户提交短信验证码（由 RPA 流程等待）。
    前端在收到 __SMS_REQUIRED__ 事件后调用此端点。
    """
    from backend.rpa.sms_bridge import submit_sms
    session_id = data.get("session_id", "")
    code = data.get("code", "").strip()
    if not session_id or not code:
        return {"success": False, "error": "参数不完整"}
    ok = submit_sms(session_id, code)
    return {"success": ok, "error": "" if ok else "会话不存在或已过期"}


@router.post("/sms/session")
async def create_sms_session():
    """创建短信验证码等待会话（前端在点击运行时调用）。"""
    from backend.rpa.sms_bridge import create_session
    session_id = create_session()
    return {"success": True, "session_id": session_id}


@router.post("/run", response_model=RpaTaskResponse)
async def run_rpa_task(
    req: RpaTaskRequest,
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    """运行 RPA 自动化任务。"""
    user = None
    if authorization:
        try:
            user = get_current_user(authorization.replace("Bearer ", ""), db)
        except Exception:
            user = None  # token 过期或无效，不阻止 RPA 执行

    from backend.rpa import run_browser_task
    result = await run_browser_task(req.task_type, req.params)

    # 异步发送通知邮件
    if user:
        task_names = {
            "port_query": "港口集装箱查询",
            "port_status": "码头状态查询",
            "track_cargo": "货物跟踪",
        }
        task_name = task_names.get(req.task_type, req.task_type)
        _send_completion_notification(user, task_name, result)

    return RpaTaskResponse(**result)


@router.post("/run/stream")
async def run_rpa_task_stream(
    req: RpaTaskRequest,
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    """SSE 流式运行 RPA 任务，实时推送日志行。"""
    user = None
    if authorization:
        try:
            user = get_current_user(authorization.replace("Bearer ", ""), db)
        except Exception:
            user = None  # token 过期或无效，不阻止 RPA 执行

    log_queue: _queue.Queue = _queue.Queue()

    async def event_stream():
        from backend.rpa import run_browser_task, set_log_queue

        # 设置全局日志队列（供 rpa_log 调用）
        set_log_queue(log_queue)

        # yield 一次以确认连接建立
        yield "event: log\ndata: [SSE connected]\n\n"

        async def _run_and_signal():
            result = await run_browser_task(req.task_type, req.params, log_queue)
            log_queue.put("__DONE__")
            return result

        task = asyncio.create_task(_run_and_signal())

        while True:
            lines = []
            while not log_queue.empty():
                lines.append(log_queue.get_nowait())

            if lines:
                sentinel_idx = None
                for i, l in enumerate(lines):
                    if l == "__DONE__":
                        sentinel_idx = i
                        continue
                    yield f"event: log\ndata: {l}\n\n"

                if sentinel_idx is not None:
                    final = await task
                    yield f"event: done\ndata: {json.dumps(final, ensure_ascii=False)}\n\n"
                    return

            await asyncio.sleep(0.3)

    return StreamingResponse(
        event_stream(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


@router.post("/letter/generate")
def generate_letter(params: dict, authorization: Optional[str] = Header(None)):
    """生成保函（非危/电放）。"""
    letter_type = params.get("type", "non_hazardous")
    carrier = params.get("carrier", "")
    data = params.get("data", {})

    if letter_type == "non_hazardous":
        content = _generate_non_hazardous_letter(carrier, data)
    elif letter_type == "telex":
        content = _generate_telex_letter(carrier, data)
    else:
        return {"success": False, "error": f"未知保函类型: {letter_type}"}

    return {"success": True, "content": content}


def _generate_non_hazardous_letter(carrier: str, data: dict) -> str:
    """生成非危保函。"""
    return f"""
NON-HAZARDOUS CARGO DECLARATION

TO: {carrier}
DATE: {data.get('date', '_______________')}

Dear Sirs,

We hereby declare that the following cargo is non-hazardous:

Shipper / 发货人: {data.get('shipper', '_______________')}
Consignee / 收货人: {data.get('consignee', '_______________')}
Port of Loading: {data.get('pol', '_______________')}
Port of Discharge: {data.get('pod', '_______________')}
Container No: {data.get('container_no', '_______________')}
Commodity: {data.get('commodity', '_______________')}
B/L No: {data.get('bl_no', '_______________')}

We confirm that the above cargo is NOT classified as dangerous goods
according to IMDG Code.

Signed by: __________________
Title: __________________
Company Seal: __________________
"""


def _generate_telex_letter(carrier: str, data: dict) -> str:
    """生成电放保函。"""
    return f"""
TELEX RELEASE LETTER

TO: {carrier}
DATE: {data.get('date', '_______________')}

Dear Sirs,

Please kindly release the cargo without presentation of original Bills of Lading.

B/L No: {data.get('bl_no', '_______________')}
Container No: {data.get('container_no', '_______________')}
Vessel/Voyage: {data.get('vessel', '_______________')}
Port of Loading: {data.get('pol', '_______________')}
Port of Discharge: {data.get('pod', '_______________')}

Shipper / 发货人: {data.get('shipper', '_______________')}
Consignee / 收货人: {data.get('consignee', '_______________')}

We hereby request you to release the cargo to the consignee
without the original Bills of Lading. All charges in connection
with this Telex Release shall be for our account.

We indemnify you against all consequences of releasing the cargo
as above requested.

Signed by: __________________
Title: __________________
Company Seal: __________________
"""


def _write_file_bytes(path: str, content: bytes) -> None:
    """Write bytes to a file (meant for asyncio.to_thread)."""
    with open(path, "wb") as f:
        f.write(content)


def _extract_bl_text(file_path: str) -> str:
    """Extract B/L text with coordinate-aware layout.

    For native PDFs: uses PyMuPDF get_text("blocks") to get text with
    #y=pos/L|C|R position tags (L=left, C=center, R=right column).

    For scanned PDFs: uses easyocr with y-position tracking.

    The position tags help the LLM understand document layout
    (which text block is a field label vs a field value, which
    column each block belongs to).
    """
    ext = os.path.splitext(file_path)[1].lower()
    if ext != ".pdf":
        return extract_text(file_path)

    import fitz
    doc = fitz.open(file_path)
    try:
        # ---- Quick check: native PDF text or scanned? ----
        has_native = False
        for pn in range(min(doc.page_count, 3)):
            blocks = doc[pn].get_text("blocks", sort=True)
            if blocks and any(b[4].strip() for b in blocks):
                has_native = True
                break

        if has_native:
            return _format_pdf_blocks(doc, file_path)
        else:
            return _ocr_pdf_blocks(doc, file_path)
    finally:
        doc.close()


def _format_pdf_blocks(doc, file_path: str) -> str:
    """PyMuPDF: extract text blocks with #y=pos/col position tags."""
    all_pages = []
    for page_num in range(doc.page_count):
        blocks = doc[page_num].get_text("blocks", sort=True)
        page_lines = []
        for b in blocks:
            text = b[4].strip()
            if not text:
                continue
            text = ' '.join(text.split())  # normalize whitespace
            y0, x0 = int(b[1]), int(b[0])
            col = "L" if x0 < 350 else ("R" if x0 > 500 else "C")
            page_lines.append(f"#y{y0}/{col} {text}")
        if page_lines:
            all_pages.append(f"--- Page {page_num + 1} ---")
            all_pages.extend(page_lines)
    result = "\n".join(all_pages)
    return result if result else extract_text(file_path)


def _ocr_pdf_blocks(doc, file_path: str) -> str:
    """OCR scanned PDF pages with easyocr, attaching y-position tags."""
    try:
        import easyocr
        reader = easyocr.Reader(["ch_sim", "en"], gpu=False)
    except ImportError:
        return extract_text(file_path)

    all_pages = []
    for page_num in range(doc.page_count):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=400)
        img_bytes = pix.tobytes("png")
        results = reader.readtext(img_bytes, paragraph=False)
        page_lines = []
        for r in results:
            text = r[1].strip()
            if not text:
                continue
            # r[0] = [[x0,y0],[x1,y0],[x1,y1],[x0,y1]]
            y_center = (r[0][0][1] + r[0][2][1]) / 2
            x_center = (r[0][0][0] + r[0][1][0]) / 2
            col = "L" if x_center < 350 else ("R" if x_center > 500 else "C")
            page_lines.append(f"#y{int(y_center)}/{col} {text}")
        if page_lines:
            all_pages.append(f"--- Page {page_num + 1} ---")
            all_pages.extend(page_lines)

    result = "\n".join(all_pages)
    return result if result else extract_text(file_path)


def _fill_telex_text_template(template_path: str, extracted: dict) -> str:
    """Fill a text-based template (.doc / .txt / .html) by simple placeholder replacement.
    Returns the filled text string (not a file path).
    """
    try:
        text = extract_text(template_path)
    except Exception:
        raise ValueError(f"Cannot read template file: {template_path}")

    vessel = extracted.get("vessel", "")
    voyage = extracted.get("voyage", "")
    vessel_voyage = f"{vessel} {voyage}".strip() if vessel or voyage else ""
    pol = extracted.get("pol", "")
    pod = extracted.get("pod", "")
    pol_pod = f"{pol} / {pod}".strip() if pol or pod else ""
    bl_no = extracted.get("bl_no", "")
    date = extracted.get("date", "")
    bl_info = ", ".join(filter(None, [bl_no, date, pol]))
    container_no = extracted.get("container_no", "")
    shipper = extracted.get("shipper", "")
    shipper_address = extracted.get("shipper_address", "")
    shipper_full = f"{shipper}\n{shipper_address}".strip() if shipper_address else shipper
    consignee = extracted.get("consignee", "")
    consignee_details = extracted.get("consignee_details", "") or consignee

    text_placeholders = {
        "[发货人抬头]": shipper_full,
        "[发货人名称]": shipper,
        "[收货人抬头]": consignee,
        "[收货人详细信息]": consignee_details,
        "[目的港]": pod,
        "[船名]": vessel,
        "[船名/航次]": vessel_voyage,
        "[起运港/目的港]": pol_pod,
        "[品名]": extracted.get("cargo_description", ""),
        "[提单号,开船日期，起运港]": bl_info,
        "[提单号, 开船日期，起运港]": bl_info,
        "[箱号]": container_no,
    }

    for ph, val in text_placeholders.items():
        if val:
            text = text.replace(ph, val)

    # Also try to replace any remaining obvious missing fields
    for ph, val in text_placeholders.items():
        text = text.replace(ph, "_______________")

    return text


def _fill_telex_docx(template_path: str, extracted: dict, carrier: str) -> str:
    """Replace placeholders in DOCX template with extracted B/L data.

    Returns path to the filled DOCX file.
    """
    import docx
    from docx import Document as DocxDocument

    doc = DocxDocument(template_path)

    # Build replacement map
    vessel = extracted.get("vessel", "")
    voyage = extracted.get("voyage", "")
    vessel_voyage = f"{vessel} {voyage}".strip() if vessel or voyage else ""
    pol = extracted.get("pol", "")
    pod = extracted.get("pod", "")
    pol_pod = f"{pol} / {pod}".strip() if pol or pod else ""
    bl_no = extracted.get("bl_no", "")
    date = extracted.get("date", "")
    bl_info = ", ".join(filter(None, [bl_no, date, pol]))
    container_no = extracted.get("container_no", "")
    shipper = extracted.get("shipper", "")
    shipper_address = extracted.get("shipper_address", "")
    shipper_full = f"{shipper}\n{shipper_address}".strip() if shipper_address else shipper
    consignee = extracted.get("consignee", "")
    consignee_details = extracted.get("consignee_details", "") or consignee

    placeholders = {
        "[发货人抬头]": shipper_full,
        "[发货人名称]": shipper,
        "[收货人抬头]": consignee,
        "[收货人详细信息]": consignee_details,
        "[目的港]": pod,
        "[船名]": vessel,
        "[船名/航次]": vessel_voyage,
        "[起运港/目的港]": pol_pod,
        "[品名]": extracted.get("cargo_description", ""),
        "[提单号,开船日期，起运港]": bl_info,
        "[提单号, 开船日期，起运港]": bl_info,
        "[箱号]": container_no,
    }

    def _replace_in_text(text: str) -> str:
        for ph, val in placeholders.items():
            if val:
                text = text.replace(ph, val)
        return text

    # Process paragraphs
    for para in doc.paragraphs:
        if not para.runs:
            continue
        # Merge all runs into first run to handle split placeholders
        full_text = "".join(r.text for r in para.runs)
        new_text = _replace_in_text(full_text)
        if new_text != full_text:
            # Clear all runs and set merged text in first run
            first = para.runs[0]
            for r in para.runs[1:]:
                r.text = ""
            first.text = new_text

    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.runs:
                        continue
                    full_text = "".join(r.text for r in para.runs)
                    new_text = _replace_in_text(full_text)
                    if new_text != full_text:
                        first = para.runs[0]
                        for r in para.runs[1:]:
                            r.text = ""
                        first.text = new_text

    # Process headers and footers
    for section in doc.sections:
        for hdr in (section.header, getattr(section, 'first_page_header', None)):
            if hdr is None:
                continue
            for para in hdr.paragraphs:
                if not para.runs:
                    continue
                full_text = "".join(r.text for r in para.runs)
                new_text = _replace_in_text(full_text)
                if new_text != full_text:
                    first = para.runs[0]
                    for r in para.runs[1:]:
                        r.text = ""
                    first.text = new_text
        for ftr in (section.footer, getattr(section, 'first_page_footer', None)):
            if ftr is None:
                continue
            for para in ftr.paragraphs:
                if not para.runs:
                    continue
                full_text = "".join(r.text for r in para.runs)
                new_text = _replace_in_text(full_text)
                if new_text != full_text:
                    first = para.runs[0]
                    for r in para.runs[1:]:
                        r.text = ""
                    first.text = new_text

    # Save to temp file
    out_dir = os.path.join(os.path.dirname(template_path), "filled")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"telex_filled_{uuid.uuid4().hex}.docx")
    doc.save(out_path)
    return out_path


def _fill_telex_doc_com(template_path: str, extracted: dict, carrier: str) -> str:
    """Use Word COM to replace placeholders in a .doc template natively.

    Spawns a subprocess to handle COM (avoids apartment-model threading issues).
    Opens the .doc in Word, does Find/Replace for each placeholder,
    saves as .docx — preserves all original formatting including headers/footers.
    """
    import json, subprocess, sys, tempfile, shutil

    # Build replacement map (same as _fill_telex_docx)
    vessel = extracted.get("vessel", "")
    voyage = extracted.get("voyage", "")
    vessel_voyage = f"{vessel} {voyage}".strip() if vessel or voyage else ""
    pol = extracted.get("pol", "")
    pod = extracted.get("pod", "")
    pol_pod = f"{pol} / {pod}".strip() if pol or pod else ""
    bl_no = extracted.get("bl_no", "")
    date = extracted.get("date", "")
    bl_info = ", ".join(filter(None, [bl_no, date, pol]))
    container_no = extracted.get("container_no", "")
    shipper = extracted.get("shipper", "")
    shipper_address = extracted.get("shipper_address", "")
    shipper_full = f"{shipper}\n{shipper_address}".strip() if shipper_address else shipper
    consignee = extracted.get("consignee", "")
    consignee_details = extracted.get("consignee_details", "") or consignee

    placeholders = {
        "[发货人抬头]": shipper_full,
        "[发货人名称]": shipper,
        "[收货人抬头]": consignee,
        "[收货人详细信息]": consignee_details,
        "[目的港]": pod,
        "[船名]": vessel,
        "[船名/航次]": vessel_voyage,
        "[起运港/目的港]": pol_pod,
        "[品名]": extracted.get("cargo_description", ""),
        "[提单号,开船日期，起运港]": bl_info,
        "[提单号, 开船日期，起运港]": bl_info,
        "[箱号]": container_no,
    }

    out_dir = os.path.join(os.path.dirname(template_path), "filled")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"telex_filled_{uuid.uuid4().hex}.docx")

    # Write a helper script that does the COM work in a dedicated process
    _script = tempfile.mktemp(suffix=".py")
    _in_path = tempfile.mktemp(suffix=".json")
    try:
        with open(_in_path, "w", encoding="utf-8") as f:
            json.dump({"input_doc": template_path, "output_docx": out_path,
                       "placeholders": placeholders}, f, ensure_ascii=False)

        with open(_script, "w", encoding="utf-8") as f:
            f.write(r'''
import json, sys, os
sys.stdout.reconfigure(encoding='utf-8')
_in = sys.argv[1]
with open(_in, "r", encoding="utf-8") as fp:
    data = json.load(fp)

import win32com.client, pythoncom
pythoncom.CoInitialize()
word = None
try:
    word = win32com.client.Dispatch("Word.Application")
    word.Visible = False
    word.DisplayAlerts = False

    import tempfile, shutil
    _tmp_fd, _tmp_path = tempfile.mkstemp(suffix=".doc")
    os.close(_tmp_fd)
    shutil.copy2(data["input_doc"], _tmp_path)
    abs_path = os.path.abspath(_tmp_path).replace("/", "\\")
    doc = word.Documents.Open(abs_path, ReadOnly=False, AddToRecentFiles=False)

    for ph, val in data["placeholders"].items():
        if not val:
            continue
        word.Selection.HomeKey(Unit=6)
        find_obj = word.Selection.Find
        find_obj.Text = ph
        find_obj.Replacement.Text = val
        find_obj.Forward = True
        find_obj.Wrap = 1
        find_obj.Format = False
        find_obj.MatchCase = True
        find_obj.MatchWholeWord = False
        find_obj.Execute(Replace=2)

    out_abs = os.path.abspath(data["output_docx"]).replace("/", "\\")
    doc.SaveAs(out_abs, FileFormat=16)
    doc.Close(SaveChanges=False)
    print("SUCCESS:" + data["output_docx"], flush=True)
except Exception as e:
    print(f"ERROR:{e}", flush=True)
    sys.exit(1)
finally:
    if word:
        try: word.Quit()
        except: pass
    pythoncom.CoUninitialize()
    try: os.unlink(_tmp_path)
    except: pass
'''.strip())

        result = subprocess.run(
            [sys.executable, _script, _in_path],
            capture_output=True, timeout=120, text=True,
            encoding="utf-8", errors="replace",
        )
        if result.returncode == 0:
            for line in result.stdout.splitlines():
                if line.startswith("SUCCESS:"):
                    return line[len("SUCCESS:"):]
        raise RuntimeError(
            f"COM subprocess failed (exit={result.returncode}): "
            f"{result.stdout[:200]} {result.stderr[:200]}"
        )
    finally:
        for _p in (_script, _in_path):
            try:
                if os.path.exists(_p):
                    os.unlink(_p)
            except Exception:
                pass


# ── Generic Phase 2: LLM-driven template filling ──

def _analyze_docx_structure(template_path: str) -> dict:
    """Extract full document structure from a DOCX for LLM analysis.

    Returns every text element with its structural position so the LLM
    can issue precise replacement instructions.
    """
    from docx import Document as DocxDocument
    doc = DocxDocument(template_path)
    structure: dict = {"paragraphs": [], "tables": [], "headers": [], "footers": []}

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            structure["paragraphs"].append({"idx": i, "text": text})

    for ti, table in enumerate(doc.tables):
        rows: list[dict] = []
        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                text = cell.text.strip()
                if text:
                    rows.append({"row": ri, "col": ci, "text": text})
        if rows:
            structure["tables"].append({"idx": ti, "rows": rows})

    for si, section in enumerate(doc.sections):
        for hdr in (section.header, getattr(section, "first_page_header", None)):
            if hdr is None:
                continue
            for pi, para in enumerate(hdr.paragraphs):
                text = para.text.strip()
                if text:
                    structure["headers"].append({"section": si, "idx": pi, "text": text})
        for ftr in (section.footer, getattr(section, "first_page_footer", None)):
            if ftr is None:
                continue
            for pi, para in enumerate(ftr.paragraphs):
                text = para.text.strip()
                if text:
                    structure["footers"].append({"section": si, "idx": pi, "text": text})

    return structure


def _apply_docx_replacements(template_path: str, out_path: str, replacements: list) -> None:
    """Apply LLM-generated {type, old, new, ...} instructions to a DOCX in-place
    so the original paragraph / table / header formatting is preserved.

    `replacements` is a Python list (already parsed from the LLM JSON reply).
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(template_path)

    def _replace_in_para(para, old_text: str, new_text: str) -> bool:
        if not para.runs:
            return False
        full = "".join(r.text for r in para.runs)
        if old_text not in full:
            return False
        new_full = full.replace(old_text, new_text)
        if new_full == full:
            return False
        first = para.runs[0]
        for r in para.runs[1:]:
            r.text = ""
        first.text = new_full
        return True

    def _replace_in_cell(cell, old_text: str, new_text: str) -> bool:
        ok = False
        for para in cell.paragraphs:
            ok = _replace_in_para(para, old_text, new_text) or ok
        return ok

    for r in replacements:
        rtype = r.get("type", "")
        old = r.get("old", "")
        new_ = r.get("new", "")
        if not old or not new_:
            continue

        if rtype == "paragraph":
            idx = r.get("idx", -1)
            if 0 <= idx < len(doc.paragraphs):
                _replace_in_para(doc.paragraphs[idx], old, new_)

        elif rtype == "table":
            ti = r.get("table_idx", -1)
            ri = r.get("row", -1)
            ci = r.get("col", -1)
            if 0 <= ti < len(doc.tables):
                tbl = doc.tables[ti]
                if 0 <= ri < len(tbl.rows) and 0 <= ci < len(tbl.rows[ri].cells):
                    _replace_in_cell(tbl.rows[ri].cells[ci], old, new_)

        elif rtype == "header":
            si = r.get("section", 0)
            idx = r.get("idx", -1)
            if 0 <= si < len(doc.sections):
                section = doc.sections[si]
                for hdr in (section.header, getattr(section, "first_page_header", None)):
                    if hdr is None:
                        continue
                    if 0 <= idx < len(hdr.paragraphs):
                        _replace_in_para(hdr.paragraphs[idx], old, new_)

        elif rtype == "footer":
            si = r.get("section", 0)
            idx = r.get("idx", -1)
            if 0 <= si < len(doc.sections):
                section = doc.sections[si]
                for ftr in (section.footer, getattr(section, "first_page_footer", None)):
                    if ftr is None:
                        continue
                    if 0 <= idx < len(ftr.paragraphs):
                        _replace_in_para(ftr.paragraphs[idx], old, new_)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)


def _apply_global_replacements(template_path: str, out_path: str, replacements: list[dict]) -> None:
    """Apply a list of {old, new} text replacements globally across ALL
    paragraphs, tables, headers and footers in a DOCX — no indices needed.
    """
    from docx import Document as DocxDocument
    doc = DocxDocument(template_path)

    def _replace_in_para(para, old: str, new: str) -> bool:
        if not para.runs:
            return False
        full = "".join(r.text for r in para.runs)
        if old not in full:
            return False
        first = para.runs[0]
        for r in para.runs[1:]:
            r.text = ""
        first.text = full.replace(old, new)
        return True

    for r in replacements:
        old = r.get("old", "")
        new_ = r.get("new", "")
        if not old or not new_:
            continue
        for para in doc.paragraphs:
            _replace_in_para(para, old, new_)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        _replace_in_para(para, old, new_)
        for section in doc.sections:
            for hdr in (section.header, getattr(section, "first_page_header", None)):
                if hdr:
                    for para in hdr.paragraphs:
                        _replace_in_para(para, old, new_)
            for ftr in (section.footer, getattr(section, "first_page_footer", None)):
                if ftr:
                    for para in ftr.paragraphs:
                        _replace_in_para(para, old, new_)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)


def _regex_gap_fill(template_path: str, out_path: str, extracted: dict) -> None:
    """Fill common placeholder patterns via regex + extracted data.
    No LLM needed — reliable even when API is down.

    Uses precise pattern matching instead of broad keyword+replaceAll approach.
    Handles: ......, …………, ____, (shipper's name), (consignee's detail),
    [insert ...], and number-labeled placeholders (⑻⑼⑽⑾⑿⒀).
    Also removes reference labels and cleans up extra whitespace.
    """
    from docx import Document as DocxDocument
    doc = DocxDocument(template_path)

    def _val(k):
        return (extracted.get(k) or "").strip()

    # ── Precise pattern matchers: (regex, replacement_or_callable) ──
    # Each pattern targets a SPECIFIC placeholder region, NOT all dots in the para.
    # Patterns are ordered: more specific before general.

    def _make_patterns(e):
        """Build list of (pattern, replacement) tuples from extracted data."""
        s = _val("shipper")
        sa = _val("shipper_address")
        c = _val("consignee")
        cd = _val("consignee_details") or c
        v = _val("vessel")
        vy = _val("voyage")
        b = _val("bl_no")
        p = _val("pol")
        pd = _val("pod")
        ct = _val("container_no")
        g = _val("cargo_description")
        d = _val("date")
        # For sentence context (Messrs), use only the first company name
        s_line = s.split('\n')[0] if s else s

        return [
            # ═══ "Attach shipper Title Paper" → shipper + address ═══
            (r'^Attach\s+shipper\s+Title\s+Paper$',
             f"{s}\n{sa}" if sa else s),

            # ═══ Labeled lines (header fields) ═══
            (r'Vessel\s*:\s*[\.…]+',
             f"Vessel : {v}" if v else None),
            (r'Voyage\s*N[°o]\s*:\s*[\.…]+',
             f"Voyage N° : {vy}" if vy else None),
            (r'Port\s+of\s+loading\s*:\s*[\.…]+',
             f"Port of loading : {p}" if p else None),
            (r'Port\s+of\s+discharge\s*:\s*[\.…]+',
             f"Port of discharge : {pd}" if pd else None),
            (r'Container\s+number\(s\)\s*:\s*[\.…]+',
             f"Container number(s) : {ct}" if ct else None),
            (r'Goods\s+description\s*:\s*[\.…]+',
             f"Goods description : {g}" if g else None),
            (r'Date\s*:\s*[\.…]+\s*',
             f"Date : {d} " if d else None),

            # ═══ B/L N° with label ⑻ ═══
            # "B/L N° : ..... ⑻..................... " → "B/L N° : SHZ8271252"
            (r'B/L\s*N[°o]\s*:\s*[\.…]+\s*[⑻]\s*[\.…]+\s*',
             f"B/L N° : {b}" if b else None),

            # ═══ Messrs pattern (complete segment with dots, shipper label, dots) ═══
            # "Messrs...(shipper's name)..... ⑼.................."
            # The entire shipper segment → "Messrs {shipper}" (one line only)
            (r'Messrs\.{2,}\([^)]*shipper[^)]*\)\s*[\.…]+\s*[⑼]+\s*[\.…]+\s*',
             f"Messrs {s_line} " if s_line else None),
            # Fallback: just (shipper's name) without Messrs prefix — one line
            (r'\([^)]*shipper[^)]*\)',
             f" {s_line} " if s_line else None),

            # ═══ Messrs consignee — labeled segment ═══
            # "Messrs....(consignee's name).............. ⑽.......... ."
            # Entire consignee segment → "Messrs {consignee}. "
            (r'Messrs\.{2,}\([^)]*consignee[^)]*name[^)]*\)\s*[\.…]+\s*[⑽]+\s*[\.…]+\s*\.?\s*',
             f"Messrs {c}. " if c else None),
            # Fallback: just (consignee's name) without Messrs
            (r'\([^)]*consignee[^)]*name[^)]*\)',
             f" {c} " if c else None),

            # ═══ (consignee's detail) in sentence context → use consignee NAME only ═══
            # (the full details block [insert ...] below gets consignee_details)
            (r'\([^)]*consignee[^)]*detail[^)]*\)',
             f" {c} " if c else None),

            # ═══ [insert ...] ═══
            (r'\[insert[^\]]*\]',
             cd if cd else None),

            # ═══ B/L N° with label ⑾ in sentence ═══
            # "N°.........⑾......." in "original bills of lading N°.........⑾....... have been surrendered"
            (r'N[°o][\.…]+[⑾][\.…]+',
             f"N° {b}" if b else None),

            # ═══ B/L N° with label ⑿ in second sentence ═══
            # "N°…………⑿……………" in "original bills of lading N°…………⑿……………..to Messrs"
            (r'N[°o][\.…]+[⑿][\.…]+\s*',
             f"N° {b} " if b else None),
        ]

    def _apply_to_para(para):
        full = "".join(r.text for r in para.runs)
        if not full:
            return
        # Skip fully-filled container numbers
        if re.search(r'\b[A-Z]{4}\d{7}\b', full):
            return
        # Skip if no placeholder patterns at all
        if not re.search(r'\.{3,}|…{2,}|_{3,}|\([^)]*name[^)]*\)|\[insert[^\]]*\]|Attach\s+shipper', full):
            return

        t = full
        patterns = _make_patterns(extracted)
        for pat, repl in patterns:
            if repl is None:
                continue
            # repl can be a string (direct replacement) or callable
            replacement = repl(extracted) if callable(repl) else repl
            try:
                t = re.sub(pat, replacement, t)
            except Exception:
                pass

        # ── Cleanup pass: remove any remaining reference labels ──
        t = re.sub(r'[⑻⑼⑽⑾⑿⒀⒁⒂⒃]', '', t)

        # ── Cleanup pass: remaining unreplaced dot groups ──
        # Only clean dots that are clearly placeholder remnants (not in long words)
        t = re.sub(r' [\.…]{4,} ', '  ', t)
        t = re.sub(r' [\.…]{4,}$', ' ', t)
        t = re.sub(r'^[\.…]{4,} ', ' ', t)

        # ── Normalize whitespace ──
        t = re.sub(r' {3,}', '  ', t)
        t = re.sub(r' ,', ',', t)
        t = re.sub(r' \.', '.', t)
        t = t.strip()

        if t == full.strip():
            return
        if para.runs:
            para.runs[0].text = t
            for r in para.runs[1:]:
                r.text = ""

    # ═══ Label cleanup pass (ALL paragraphs, not just those with placeholders) ═══
    # Remove reference markers: circled numbers and (14)(15)(16) style
    _label_pattern = re.compile(r'[①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃]|\(\d{2}\)')

    def _remove_numbering(para):
        """Strip Word auto-numbering (w:numPr) from paragraph XML.
        This handles bullet/numbering that python-docx text can't see."""
        import lxml.etree as _lxml_etree
        ns = 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'
        ppr = para._element.find(f'{{{ns}}}pPr')
        if ppr is not None:
            numPr = ppr.find(f'{{{ns}}}numPr')
            if numPr is not None:
                ppr.remove(numPr)

    def _cleanup_labels(para):
        full = "".join(r.text for r in para.runs)
        if not full:
            return
        new = _label_pattern.sub('', full)
        new = re.sub(r' {2,}', ' ', new).strip()
        if new != full:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""

    for para in doc.paragraphs:
        _remove_numbering(para)
        _apply_to_para(para)
        _cleanup_labels(para)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _remove_numbering(para)
                    _apply_to_para(para)
                    _cleanup_labels(para)
    for section in doc.sections:
        for hdr in (section.header, getattr(section, "first_page_header", None)):
            if hdr:
                for para in hdr.paragraphs:
                    _remove_numbering(para)
                    _apply_to_para(para)
                    _cleanup_labels(para)
        for ftr in (section.footer, getattr(section, "first_page_footer", None)):
            if ftr:
                for para in ftr.paragraphs:
                    _remove_numbering(para)
                    _apply_to_para(para)
                    _cleanup_labels(para)

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)


def _fill_nonhazardous_docx(template_path: str, extracted: dict, carrier: str) -> str:
    """Replace field blanks in non-hazardous letter DOCX template with extracted data.

    Handles:
    - Paragraphs with field label + colon + underscore blanks (例1/例2)
    - Table rows with field label in one cell and blank in adjacent cell (例3)

    Returns path to the filled DOCX file.
    """
    from docx import Document as DocxDocument

    doc = DocxDocument(template_path)

    def _val(k):
        v = extracted.get(k, "")
        return v.strip() if v else ""

    # Build field→value mapping
    fv = {
        "vv": _val("vessel_voyage") or f"{_val('vessel')} {_val('voyage')}".strip(),
        "bl": _val("bl_no"),
        "cntr": _val("container_no"),
        "pol": _val("pol"),
        "pod": _val("pod"),
        "cn_cn": _val("中文品名"),
        "en_cn": _val("英文品名"),
        "cas": _val("CAS_NO"),
        "app": _val("外观与性状"),
        "use": _val("主要用途"),
        "shp": _val("shipper"),
        "cne": _val("consignee"),
    }

    # (label_keywords, field_key) — ordered by specificity
    label_map = [
        (["船名/航次", "VSL/VOY", "VSL", "VOY", "Vessel", "Voyage"], "vv"),
        (["提单号", "BL No", "BL NO", "B/L No", "B/L NO", "关单号码"], "bl"),
        (["箱型/箱量", "箱号", "CNTR NO", "Container Type", "Container", "CNTR"], "cntr"),
        (["起运港", "Port of Loading", "POL"], "pol"),
        (["目的港", "卸货港", "Port of Discharge", "Final destination", "POD"], "pod"),
        (["中文品名", "Commodity description in Chinese"], "cn_cn"),
        (["英文品名", "Commodity description in English", "Commodity Name"], "en_cn"),
        (["CAS NO", "CAS号", "CAS number", "化学品文摘号"], "cas"),
        (["外观与性状", "Commodity shape and property", "Appearance", "样品外观"], "app"),
        (["主要用途", "Utility of Goods", "Cargo Usage", "用途", "Usage"], "use"),
        (["发货人", "Shipper", "货代"], "shp"),
        (["收货人", "Consignee"], "cne"),
        (["公司名", "Customer Name"], "shp"),
    ]

    def _replace_para(text):
        """If text contains a known field label, replace blank/underscore after colon with value."""
        if not text:
            return text
        fk = None
        for labels, key in label_map:
            for lbl in labels:
                if lbl in text:
                    fk = key
                    break
            if fk:
                break
        if fk is None:
            return text
        val = fv.get(fk, "")
        if not val:
            return text
        # Find the first colon and replace content after it
        cm = re.search(r'[：:]\s*', text)
        if not cm:
            return text
        left = text[:cm.end()]
        right = text[cm.end():]
        # Check if right side is mostly blank/underscore
        non_blank = re.sub(r'[\s_]+', '', right)
        if len(non_blank) <= 2:
            return f"{left} {val}"
        # Otherwise try replacing longest underscore block
        blocks = re.findall(r'_{3,}', right)
        if blocks:
            longest = max(blocks, key=len)
            return text.replace(longest, val, 1)
        return text

    # ── Paragraphs ──
    for para in doc.paragraphs:
        if not para.runs:
            continue
        full = "".join(r.text for r in para.runs)
        new = _replace_para(full)
        if new != full:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""

    # ── Tables ──
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) < 2:
                continue
            fk = None
            for labels, key in label_map:
                for lbl in labels:
                    if lbl in cells[0]:
                        fk = key
                        break
                if fk:
                    break
            if fk is None:
                continue
            val = fv.get(fk, "")
            if not val:
                continue
            # Fill second cell
            for para in row.cells[1].paragraphs:
                if para.runs:
                    para.runs[0].text = val
                    for r in para.runs[1:]:
                        r.text = ""
                else:
                    para.add_run(val)

    out_dir = os.path.join(os.path.dirname(template_path), "filled")
    os.makedirs(out_dir, exist_ok=True)
    out_path = os.path.join(out_dir, f"nonhazardous_filled_{uuid.uuid4().hex}.docx")
    doc.save(out_path)
    return out_path


def _regex_gap_fill_nonhazardous(template_path: str, out_path: str, extracted: dict) -> None:
    """Clean up remaining blanks in non-hazardous letter after hardcoded fill."""
    from docx import Document as DocxDocument
    doc = DocxDocument(template_path)

    def _apply(text):
        if not text:
            return text
        text = re.sub(r'____年____月____日', '______年______月______日', text)
        text = re.sub(r'_{10,}', '_______________', text)
        return text

    for para in doc.paragraphs:
        if not para.runs:
            continue
        full = "".join(r.text for r in para.runs)
        new = _apply(full)
        if new != full:
            para.runs[0].text = new
            for r in para.runs[1:]:
                r.text = ""

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if not para.runs:
                        continue
                    full = "".join(r.text for r in para.runs)
                    new = _apply(full)
                    if new != full:
                        para.runs[0].text = new
                        for r in para.runs[1:]:
                            r.text = ""

    os.makedirs(os.path.dirname(out_path) or ".", exist_ok=True)
    doc.save(out_path)


@router.post("/letter/auto-fill")
async def auto_fill_letter(
    msds: UploadFile = File(..., description="MSDS文档(PDF)"),
    certificate: UploadFile = File(..., description="鉴定书(PDF)"),
    template: UploadFile = File(None, description="非危保函模板(DOCX/DOC/PDF)"),
    carrier: str = Form(""),
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    """AI自动填写非危保函 - 解析MSDS和鉴定书，自动填充保函内容。"""
    # 1. Save uploaded files temporarily
    upload_dir = os.path.join(settings.UPLOAD_DIR, "auto_fill")
    os.makedirs(upload_dir, exist_ok=True)

    saved_paths = {}
    file_handles = {
        "msds": msds,
        "certificate": certificate,
    }
    if template:
        file_handles["template"] = template

    try:
        for name, f in file_handles.items():
            ext = os.path.splitext(f.filename or "")[1] or ".pdf"
            save_path = os.path.join(upload_dir, f"{name}_{uuid.uuid4().hex}{ext}")
            content = await f.read()
            await asyncio.to_thread(_write_file_bytes, save_path, content)
            saved_paths[name] = save_path

        # 2. Parse each file
        msds_text = await asyncio.to_thread(extract_text, saved_paths["msds"])
        cert_text = await asyncio.to_thread(extract_text, saved_paths["certificate"])

        template_text = ""
        if "template" in saved_paths:
            template_text = await asyncio.to_thread(
                extract_text, saved_paths["template"]
            )

        # 3. If no template uploaded, generate one from carrier name
        if not template_text:
            template_text = _generate_non_hazardous_letter(
                carrier or "CARRIER", {}
            )

        # 4. Build LLM prompts — extract-first approach for reliability
        # Phase 1: Extract structured data from documents
        phase1_system = (
            "You are a document extraction specialist. "
            "Extract information from the MSDS (Material Safety Data Sheet) and "
            "鉴定书 (appraisal certificate) to fill a non-hazardous cargo letter.\n\n"
            "Return ONLY a valid JSON object. No other text.\n"
            '{"中文品名":"","英文品名":"","CAS_NO":"","外观与性状":"","主要用途":"",'
            '"shipper":"","consignee":"","container_no":"","bl_no":"",'
            '"vessel_voyage":"","pol":"","pod":""}'
        )
        # Phase 1 user: truncate but be generous
        msds_section = msds_text[:8000] if len(msds_text) > 8000 else msds_text
        cert_section = cert_text[:6000] if len(cert_text) > 6000 else cert_text
        phase1_user = (
            f"## MSDS (SDS) Text:\n{msds_section}\n\n"
            f"## 鉴定书 (Appraisal Certificate) Text:\n{cert_section}\n\n"
            "Extract the following information from the documents above "
            "to fill a non-hazardous cargo declaration letter:\n"
            "- 中文品名 (Commodity description in Chinese)\n"
            "- 英文品名 (English commodity name)\n"
            "- CAS_NO (CAS registry number — look in the MSDS Ingredients table)\n"
            "- 外观与性状 (Appearance and properties / 样品外观与气味)\n"
            "- 主要用途 (Main utility of goods / 产品主要用途)\n"
            "- shipper (货主/发货人 / 委托单位 / Applicant)\n"
            "- consignee (收货人)\n"
            "- container_no (柜号/箱号)\n"
            "- bl_no (提单号/BL number)\n"
            "- vessel_voyage (船名/航次)\n"
            "- pol (起运港/port of loading)\n"
            "- pod (目的港/port of discharge)\n\n"
            "Return ONLY the JSON. Values not found → use empty string."
        )

        # Phase 2: Fill the template with extracted data
        phase2_system = (
            "You are a document filing specialist. Your task is to fill in blanks.\n\n"
            "RULES (follow exactly):\n"
            "1. Replace _______________ blanks with the extracted data\n"
            "2. If a value is empty, keep the blank as-is\n"
            "3. Do NOT add comments, reasoning, or explanations\n"
            "4. Do NOT describe what you are doing\n"
            "5. Return ONLY the filled document - no other text at all\n"
            "6. Keep the original formatting and structure of the template"
        )

        # 5. Use NVIDIA NIM for document extraction
        from backend.addons.llm.multi_agent import NVIDIA_API_BASE

        nim_key = "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"
        model_used = "nvidia/nemotron-3-super-120b-a12b"
        api_url = NVIDIA_API_BASE
        api_key = nim_key

        async def _llm_call(messages: list, max_tokens: int = 2048, temp: float = 0.05) -> str:
            """Call NVIDIA NIM with retry. Returns response text."""
            import asyncio as _asyncio

            headers = {"Content-Type": "application/json"}
            if api_key:
                headers["Authorization"] = f"Bearer {api_key}"

            url = f"{api_url.rstrip('/')}/chat/completions"
            payload = {
                "model": model_used,
                "messages": messages,
                "temperature": temp,
                "max_tokens": max_tokens,
            }

            last_err = None
            for attempt in range(3):
                try:
                    async with httpx.AsyncClient(timeout=180.0) as cli:
                        resp = await cli.post(url, headers=headers, json=payload)
                    if resp.status_code == 429:  # rate limit
                        await _asyncio.sleep(2 ** attempt)
                        continue
                    resp.raise_for_status()
                    return resp.json()["choices"][0]["message"]["content"].strip()
                except (httpx.RemoteProtocolError, httpx.ConnectError, httpx.TimeoutException) as e:
                    last_err = e
                    wait = 2 ** attempt
                    print(f"[auto-fill] API call attempt {attempt+1} failed: {e}, retry in {wait}s")
                    await _asyncio.sleep(wait)
            raise last_err or RuntimeError("LLM call failed after retries")

        # Phase 1: extract JSON from documents
        print("[auto-fill] Phase 1: extracting JSON from documents...")
        phase1_messages = [
            {"role": "system", "content": phase1_system},
            {"role": "user", "content": phase1_user},
        ]
        extracted_text = await _llm_call(phase1_messages, max_tokens=2048)
        print(f"[auto-fill] Phase 1 response: {extracted_text[:300]}")

        # ── Phase 2: DOCX pipeline (if template is DOCX/DOC) ──
        filled_docx_path = ""
        download_id = ""
        filled_letter = ""

        # ═══ Parse extracted JSON from Phase 1 response ═══
        # (moved before DOCX pipeline so it's defined when needed)
        extracted = {
            "中文品名": "",
            "英文品名": "",
            "CAS_NO": "",
            "外观与性状": "",
            "主要用途": "",
            "shipper": "",
            "consignee": "",
            "container_no": "",
            "bl_no": "",
            "vessel_voyage": "",
            "pol": "",
            "pod": "",
        }

        json_match = re.search(r"\{.*\}", extracted_text, re.DOTALL)
        if json_match:
            try:
                parsed = json.loads(json_match.group())
                if isinstance(parsed, dict):
                    extracted.update({k: v for k, v in parsed.items() if k in extracted})
            except (json.JSONDecodeError, ValueError):
                pass

        if "template" in saved_paths:
            template_path = saved_paths["template"]
            ext = os.path.splitext(template_path)[1].lower()

            if ext in (".docx", ".doc"):
                print("[auto-fill] Phase 2: DOCX pipeline...")
                try:
                    import subprocess as _sp
                    docx_path = template_path
                    if ext == ".doc":
                        converted = os.path.join(
                            upload_dir, f"convert_{uuid.uuid4().hex}.docx"
                        )
                        _sp.run(
                            ["soffice", "--headless", "--convert-to", "docx",
                             "--outdir", upload_dir, template_path],
                            capture_output=True, timeout=60, text=True
                        )
                        for f in os.listdir(upload_dir):
                            if f.endswith(".docx") and "convert_" in f:
                                docx_path = os.path.join(upload_dir, f)
                                break

                    filled_docx_path = await asyncio.to_thread(
                        _fill_nonhazardous_docx, docx_path, extracted, carrier or ""
                    )
                    # Regex gap-fill cleanup
                    _gp_dir = os.path.join(upload_dir, "gap_filled")
                    os.makedirs(_gp_dir, exist_ok=True)
                    _gp_path = os.path.join(
                        _gp_dir, f"nonhazardous_filled_{uuid.uuid4().hex}.docx"
                    )
                    _regex_gap_fill_nonhazardous(
                        filled_docx_path, _gp_path, extracted
                    )

                    filled_docx_path = _gp_path
                    download_id = str(uuid.uuid4())
                    _telex_download_store[download_id] = filled_docx_path

                    filled_letter = await asyncio.to_thread(
                        extract_text, filled_docx_path
                    )
                    print(f"  DOCX fill OK -> {os.path.basename(filled_docx_path)}")
                except Exception as _de:
                    print(f"  DOCX pipeline failed: {_de}, falling back to LLM fill")
                    filled_docx_path = ""
                    download_id = ""

        # ── Phase 2: LLM fill (fallback for PDF/no template) ──
        if not filled_letter:
            print("[auto-fill] Phase 2: LLM filling template...")
            phase2_user = (
                f"## Extracted Information (JSON):\n{extracted_text}\n\n"
                f"## Template to Fill:\n{template_text}\n\n"
                f"Carrier (船司): {carrier or 'Unknown Carrier'}\n\n"
                "Fill the template using the extracted information."
            )
            phase2_messages = [
                {"role": "system", "content": phase2_system},
                {"role": "user", "content": phase2_user},
            ]
            filled_letter = await _llm_call(phase2_messages, max_tokens=4096)

            # Clean up any markers
            filled_letter = re.sub(
                r'\[MISSING:[^\]]*\]', '_______________', filled_letter
            )
            filled_letter = re.sub(
                r"^={3,}.*?={3,}\s*", "", filled_letter, flags=re.MULTILINE
            ).strip()

            # Post-process: strip LLM preamble, find letter start
            _known_headers = [
                "非危化工品保函", "非危保函", "NON-HAZARDOUS", "Non-Hazardous",
                "Letter of Indemnity", "DG/NON-DG", "货物性质咨询申请表",
                "致森罗商船", "SHAM", "SM LINE",
            ]
            _best = filled_letter
            for h in _known_headers:
                idx = filled_letter.rfind(h)
                if idx >= 0 and idx < len(filled_letter) - 50:
                    _best = filled_letter[idx:]
                    break
            filled_letter = _best

        # Fallback: if CAS_NO not found by LLM, try regex from raw OCR text
        if not extracted.get("CAS_NO"):
            cas_match = re.search(r"\b(\d{2,7}-\d{2}-\d)\b", msds_text + "\n" + cert_text)
            if cas_match:
                extracted["CAS_NO"] = cas_match.group(1)

        return {
            "success": True,
            "filled_letter": filled_letter,
            "extracted": extracted,
            "model_used": model_used,
            "download_id": download_id,
        }

    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        print(f"[auto-fill] ERROR: {error_detail}")
        # Write to log file for debugging
        log_path = os.path.join(settings.UPLOAD_DIR, "auto_fill_errors.log")
        try:
            with open(log_path, "a", encoding="utf-8") as logf:
                logf.write(f"\n[ERROR at {__import__('datetime').datetime.now()}]\n{error_detail}\n")
        except Exception:
            pass
        return {
            "success": False,
            "error": str(e),
            "filled_letter": "",
            "extracted": {},
            "model_used": "",
            "download_id": "",
        }

    finally:
        # Clean up temp files
        for path in saved_paths.values():
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception:
                pass


@router.post("/letter/auto-fill-telex")
async def auto_fill_telex_letter(
    bill_of_lading: UploadFile = File(..., description="提单PDF"),
    template: UploadFile = File(None, description="电放保函模板(DOCX/DOC/PDF)"),
    carrier: str = Form(""),
    authorization: Optional[str] = Header(None),
    db: Session = Depends(get_db),
):
    """AI自动填充电放保函 - 解析提单内容，自动填充到电放保函模板。"""
    upload_dir = os.path.join(settings.UPLOAD_DIR, "auto_fill_telex")
    os.makedirs(upload_dir, exist_ok=True)

    saved_paths = {}
    file_handles = {"bl": bill_of_lading}
    if template:
        file_handles["template"] = template

    try:
        for name, f in file_handles.items():
            ext = os.path.splitext(f.filename or "")[1] or ".pdf"
            save_path = os.path.join(upload_dir, f"{name}_{uuid.uuid4().hex}{ext}")
            content = await f.read()
            await asyncio.to_thread(_write_file_bytes, save_path, content)
            saved_paths[name] = save_path

        # 1. Parse B/L text — use easyocr for better Chinese B/L OCR quality
        bl_text = await asyncio.to_thread(_extract_bl_text, saved_paths["bl"])

        # 2. Parse template text (if provided)
        template_text = ""
        if "template" in saved_paths:
            template_text = await asyncio.to_thread(extract_text, saved_paths["template"])

        # 3. If no template uploaded, use a default telex format
        if not template_text:
            template_text = _generate_telex_letter(carrier or "CARRIER", {})

        # 4. Phase 1: Extract structured data from B/L
        phase1_system = (
            "You are a bill of lading extraction specialist. "
            "Extract shipping information from the Bill of Lading text below. "
            "The text contains #y=pos/L|C|R tags to indicate each "
            "text block's position on the page (L=left column, C=center, R=right).\n"
            "Use these position hints:\n"
            "- Text blocks at the TOP of the page (low y values like y72-200) "
            "are usually SHIPPER / B/L NO / VESSEL info\n"
            "- Left-column blocks #.../L near the top are typically SHIPPER and CONSIGNEE\n"
            "- Right-column blocks #.../R near the top are typically B/L number and vessel info\n"
            "- Blocks with similar y-ranges belong to the same section\n"
            "- The text below a SHIPPER label at similar y-range is the shipper value\n"
            "- Table-like data (containers, weights) is often in the middle column #.../C\n\n"
            "CRITICAL: The #y.../L|C|R tags are position hints only, NEVER include "
            "them in extracted field values. Extract only the actual document text.\n\n"
            "The text may contain OCR errors — use context to correct obvious mistakes.\n\n"
            "Return ONLY a valid JSON object. No other text. No reasoning.\n"
            '{"bl_no":"","vessel":"","voyage":"","pol":"","pod":"",'
            '"shipper":"","consignee":"","container_no":"",'
            '"cargo_description":"","date":"","place_of_issue":"",'
            '"shipper_address":"","consignee_details":""}'
        )
        bl_section = bl_text[:15000] if len(bl_text) > 15000 else bl_text
        phase1_user = (
            f"## Bill of Lading Text with Position Tags #y=pos/L|C|R:\n"
            f"(Use the y-position and L/C/R column tags to understand layout)\n"
            f"{bl_section}\n\n"
            "Extract the following information from the Bill of Lading above. "
            "The text comes from OCR so some characters may be wrong — use context clues:\n"
            "- bl_no (提单号/Bill of Lading number, format example: MEDUXZ341199 or SHZ8271252)\n"
            "- vessel (船名/vessel name — some B/Ls put it near 'VESSEL' label, "
            "others only at bottom near 'SIGNED FOR THE CARRIER' / 'as agents for the carrier'. "
            "Look for the distinctive vessel name, not just the word after 'VESSEL'.)\n"
            "- voyage (航次/voyage number, often after vessel name or near 'VOYAGE')\n"
            "- pol (起运港/port of loading, near 'PORT OF LOADING')\n"
            "- pod (目的港/port of discharge, near 'PORT OF DISCHARGE')\n"
            "- shipper (发货人/shipper): The FULL shipper identification block — usually "
            "2-5 lines starting after the 'SHIPPER' / '发货人' header at the very top of the B/L.\n"
            "  CRITICAL: The shipper block may span MULTIPLE LINES. It often contains:\n"
            "    Line 1: Company name (e.g. 'FOSHAN CITY JIN LI')\n"
            "    Line 2: Trading/DBA name (e.g. 'LONGTANG TECHNOLOGY CO., LTD')\n"
            "    Line 3-N: Street address, city, postal code, country\n"
            "  Extract ALL company/trading names as the 'shipper' field — combine them "
            "with a newline if there are multiple distinct names.\n"
            "  The OCR may split 'CO., LTD' or 'CO. LTD' across two lines. "
            "Check the next line for continuation ('LTD', 'LIMITED', 'INC', 'GMBH') "
            "and INCLUDE it.\n"
            "  If you see both a Chinese and an English name, use the English name.\n"
            "  Do NOT truncate — capture the complete company identification.\n"
            "- shipper_address: The shipper's PHYSICAL ADDRESS ONLY (street, city/district, "
            "province/country, postal code).\n"
            "  CRITICAL: This field must contain ONLY address lines, NOT company names. "
            "If 'LONGTANG TECHNOLOGY CO., LTD' appears in the address block, it is a "
            "trading name, NOT an address — put it in 'shipper' instead.\n"
            "  On MSC B/Ls the address may appear near the shipper name but mixed "
            "with clause text (like 'XNQIAO VILLAGE'). Look for address-like text near "
            "the shipper name or in the area after 'SHIPPER'S LOAD STOW AND COUNT'.\n"
            "- consignee (收货人/consignee): The party receiving the goods.\n"
            "  Located below the shipper section, after the 'CONSIGNEE' header or '收货人' label.\n"
            "  On ocean B/Ls this is OFTEN 'TO ORDER' or 'TO THE ORDER OF ...'.\n"
            "  Extract the full consignee name (company name or 'TO ORDER').\n"
            "  CRITICAL: Do NOT use the shipper name as the consignee. "
            "If no clear 'CONSIGNEE' label exists in the OCR text, look for a company name "
            "block between the SHIPPER section and the NOTIFY PARTIES section.\n"
            "  If the B/L says 'TO ORDER' or 'TO THE ORDER OF', the actual consignee "
            "may appear near the NOTIFY PARTIES area — extract that company name instead.\n"
            "  If you see both Chinese and English, use the English name.\n"
            "- consignee_details: The consignee's full details including company name, "
            "full address, VAT number, phone number, and contact person. "
            "For 'TO ORDER' bills, look for the actual party receiving the goods "
            "near the NOTIFY PARTIES section. Combine all available info into one block.\n"
            "  CRITICAL: Extract ONLY the text that actually appears in the B/L. "
            "Do NOT infer, guess, or add any information that is not explicitly written. "
            "Do NOT add country names, area codes, or company names that are not present. "
            "IMPORTANT: The B/L text contains #y.../L|C|R position tags - these are "
            "layout hints, NOT part of the document content. NEVER include these tags "
            "in your extracted values. Extract only the actual document text.\n"
            "- container_no (柜号/container number, format: 4 letters + 7 digits like FFAU2256329 or KBRU2611456. "
            "Do NOT confuse with B/L number! Container numbers always follow the pattern AAAA1234567.)\n"
            "- cargo_description (品名/goods description)\n"
            "- date (装船日/shipped on board date)\n"
            "- place_of_issue (签发地/place of issue)\n\n"
            "Return ONLY the JSON. Values not found → use empty string."
        )

        # 5. Use NVIDIA NIM for document extraction (Phase 1 only)
        from backend.addons.llm.multi_agent import NVIDIA_API_BASE

        nim_key = "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"
        model_primary = "nvidia/nemotron-3-super-120b-a12b"
        model_fallback = "openai/gpt-oss-120b"
        api_url = NVIDIA_API_BASE
        api_key = nim_key

        async def _llm_call(messages: list, max_tokens: int = 4096, temp: float = 0.05) -> str:
            """Call NVIDIA NIM with retry and model fallback. Returns response text."""
            import asyncio as _asyncio
            models_to_try = [model_primary, model_fallback]
            for model_idx, model_name in enumerate(models_to_try):
                for attempt in range(3):
                    try:
                        headers = {"Content-Type": "application/json"}
                        if api_key:
                            headers["Authorization"] = f"Bearer {api_key}"
                        url = f"{api_url.rstrip('/')}/chat/completions"
                        payload = {
                            "model": model_name,
                            "messages": messages,
                            "temperature": temp,
                            "max_tokens": max_tokens,
                        }
                        async with httpx.AsyncClient(timeout=180.0) as cli:
                            resp = await cli.post(url, headers=headers, json=payload)
                        if resp.status_code in (429, 503, 502, 504):
                            wait = 2 ** attempt
                            print(f"[auto-fill-telex] {model_name} attempt {attempt+1} "
                                  f"got {resp.status_code}, retry in {wait}s")
                            await _asyncio.sleep(wait)
                            continue
                        resp.raise_for_status()
                        text = resp.json()["choices"][0]["message"]["content"].strip()
                        if model_idx == 0 and not re.search(r'\{.*\}', text, re.DOTALL):
                            print(f"[auto-fill-telex] {model_name} returned reasoning, trying fallback")
                            break
                        return text
                    except (httpx.RemoteProtocolError, httpx.ConnectError,
                            httpx.TimeoutException, httpx.HTTPStatusError) as e:
                        wait = 2 ** attempt
                        print(f"[auto-fill-telex] {model_name} attempt {attempt+1} failed: {e}, retry in {wait}s")
                        await _asyncio.sleep(wait)
            raise RuntimeError("LLM call failed after all retries")

        # Phase 1: extract JSON from B/L
        print("[auto-fill-telex] Phase 1: extracting JSON from B/L...")
        phase1_messages = [
            {"role": "system", "content": phase1_system},
            {"role": "user", "content": phase1_user},
        ]
        extracted_text = await _llm_call(phase1_messages, max_tokens=2048)
        print(f"[auto-fill-telex] Phase 1 response: {extracted_text[:300]}")

        # Parse extracted JSON
        extracted = {
            "bl_no": "", "vessel": "", "voyage": "", "pol": "", "pod": "",
            "shipper": "", "consignee": "", "container_no": "",
            "cargo_description": "", "date": "", "place_of_issue": "",
            "shipper_address": "", "consignee_details": "",
        }
        json_match = re.search(r"\{.*\}", extracted_text, re.DOTALL)
        if json_match:
            try:
                parsed = json.loads(json_match.group())
                if isinstance(parsed, dict):
                    extracted.update({k: v for k, v in parsed.items() if k in extracted})
            except (json.JSONDecodeError, ValueError):
                pass

        # Post-process: fix swapped BL no / container no
        cntr_pat = re.compile(r'^[A-Z]{4}\d{7}$')
        bl_raw = extracted.get("bl_no", "").strip().upper()
        cntr_raw = extracted.get("container_no", "").strip().upper()
        if cntr_pat.match(bl_raw) and not cntr_pat.match(cntr_raw):
            extracted["bl_no"], extracted["container_no"] = cntr_raw, bl_raw

        # Build vessel/voyage combined field
        if extracted.get("vessel") and extracted.get("voyage"):
            extracted["vessel_voyage"] = f"{extracted['vessel']} {extracted['voyage']}"
        else:
            extracted["vessel_voyage"] = extracted.get("vessel") or extracted.get("voyage") or ""

        # Fallback: always run regex as a cross-check for shipper / consignee
        bl_upper = bl_text.upper()
        bl_lines = bl_text.split("\n")

        # --- Shipper name continuation + address extraction (always run) ---
        _shipper_fallback_name = None
        _shipper_fallback_addr = None
        for sep in ["SHIPPER", "发货人", "发货人 :", "SHIPPER :"]:
            for i, line in enumerate(bl_lines):
                if sep in line.upper():
                    for j in range(i + 1, min(i + 5, len(bl_lines))):
                        candidate = bl_lines[j].strip()
                        candidate = re.sub(r'^[:\s]*|[:\s]*$', '', candidate)
                        candidate = re.sub(r'#y\d+/[LCR]\s*|\[y=\d+\|[LCR]\]\s*', '', candidate).strip()
                        if not candidate or candidate.startswith("CONSIGNEE") or candidate.startswith("收货人"):
                            continue
                        # Skip boilerplate lines that aren't company names
                        if any(kw in candidate.upper() for kw in
                               ["ENDORSEMENT", "AGENT", "LOAD", "STOW", "COUNT",
                                "SEALED", "CARRIER", "VESSEL", "PORT OF"]):
                            continue
                        # Build shipper name: check next lines for continuation
                        name_parts = [candidate]
                        continue_idx = j
                        for k in range(j + 1, min(j + 3, len(bl_lines))):
                            cont = bl_lines[k].strip()
                            cont_clean = re.sub(r'^[:\s]*|[:\s]*$', '', cont)
                            cont_clean = re.sub(r'#y\d+/[LCR]\s*|\[y=\d+\|[LCR]\]\s*', '', cont_clean).strip()
                            if not cont_clean:
                                break
                            if re.match(r'^[,\.]\s*(LTD|LIMITED|INC|GMBH|CO)\b', cont_clean, re.I) or \
                               re.match(r'^(LTD|LIMITED|INC|GMBH)\b', cont_clean, re.I):
                                name_parts.append(cont_clean)
                                continue_idx = k
                            elif re.match(r'^[A-Z]', cont_clean) and len(cont_clean) < 15:
                                name_parts.append(cont_clean)
                                continue_idx = k
                            else:
                                break
                        _shipper_fallback_name = " ".join(name_parts)
                        # Collect address lines after shipper name
                        addr_lines = []
                        for k in range(continue_idx + 1, min(continue_idx + 6, len(bl_lines))):
                            addr_line = bl_lines[k].strip()
                            addr_line = re.sub(r'^[:\s]*|[:\s]*$', '', addr_line)
                            addr_line = re.sub(r'#y\d+/[LCR]\s*|\[y=\d+\|[LCR]\]\s*', '', addr_line).strip()
                            if not addr_line:
                                break
                            if any(kw in addr_line.upper() for kw in
                                   ["CONSIGNEE", "NOTIFY PARTY", "NOTIFY PARTIE", "NOTIFY PARTIES", "收货人", "SHIPPER", "VESSEL"]):
                                break
                            if re.search(r'[A-Za-z0-9]{3,}', addr_line):
                                addr_lines.append(addr_line)
                        if addr_lines:
                            _shipper_fallback_addr = "\n".join(addr_lines)
                        break
                if _shipper_fallback_name:
                    break
            if _shipper_fallback_name:
                break
        # Apply shipper fallback (enrich AI result)
        # Only overwrite AI if fallback contains a valid legal suffix that AI misses
        if _shipper_fallback_name:
            ai_shipper = (extracted.get("shipper") or "").strip()
            if not ai_shipper or len(ai_shipper) < 5:
                extracted["shipper"] = _shipper_fallback_name
            else:
                ai_has_suffix = bool(re.search(r'\b(LTD|LIMITED|INC|GMBH)\b', ai_shipper.upper()))
                fb_has_suffix = bool(re.search(r'\b(LTD|LIMITED|INC|GMBH)\b', _shipper_fallback_name.upper()))
                if not ai_has_suffix and fb_has_suffix:
                    # Use fallback name but clean it: keep up to the legal suffix
                    clean_parts = []
                    for part in _shipper_fallback_name.split():
                        clean_parts.append(part)
                        if re.match(r'(LTD|LIMITED|INC|GMBH)[.,]*', part.upper()):
                            break
                    extracted["shipper"] = " ".join(clean_parts)

        # Post-process: if shipper ends with "CO" / "CO." and has no suffix, append "., LTD"
        _shp = extracted.get("shipper", "").strip()
        if _shp and len(_shp) >= 3:
            _shp_upper = _shp.upper()
            if not re.search(r'\b(LTD|LIMITED|INC|GMBH)\b', _shp_upper):
                if re.search(r'\bCO\.?,?\s*$', _shp_upper) or _shp_upper.endswith(" CO"):
                    _shp = _shp.rstrip('., ')
                    extracted["shipper"] = _shp + "., LTD"
                    print(f"[Shipper LTD fix] {_shp} (no suffix) → {extracted['shipper']}")

        # Shipper address: search near actual shipper name in OCR lines
        if not extracted.get("shipper_address"):
            shipper_key = (extracted.get("shipper", "").replace("CO., LTD", "")
                           .replace("CO. LTD", "").replace(", LTD", "").replace(" LTD", "")
                           .strip()[:25])
            if shipper_key:
                for i, line in enumerate(bl_lines):
                    if shipper_key in line:
                        addr = []
                        for j in range(i + 1, min(i + 12, len(bl_lines))):
                            al = bl_lines[j].strip()
                            al = re.sub(r'#y\d+/[LCR]\s*|\[y=\d+\|[LCR]\]\s*', '', al).strip()
                            if not al or len(al) < 5:
                                continue
                            if any(kw in al.upper() for kw in
                                   ["NOTIFY", "CONSIGNEE", "VESSEL", "PORT OF",
                                    "FREIGHT", "CONTAINER", "DESCRIPTION",
                                    "LLOYDS", "ISPM"]):
                                break
                            if re.search(r'(VILLAGE|STREET|ROAD|DISTRICT|PROVINCE|'
                                         r'CITY|ZONE|COUNTY|BUILDING|FLOOR|'
                                         r'NO\.?\s*\d|\d{3,}\s+[A-Z])', al, re.I):
                                addr.append(al)
                        if addr:
                            extracted["shipper_address"] = "\n".join(addr)
                        break

        # --- Consignee details extraction (always run, for consignee_details) ---
        # The AI may return the correct consignee name. This fallback enriches with
        # address/contact info but only overwrites the name if AI returned shipper.
        _consignee_parts = None

        # Try 1: find "CONSIGNEE" / "收货人" label
        for sep in ["CONSIGNEE", "收货人", "收货人 :", "CONSIGNEE :"]:
            for i, line in enumerate(bl_lines):
                if sep in line.upper():
                    parts = []
                    for j in range(i + 1, min(i + 8, len(bl_lines))):
                        candidate = bl_lines[j].strip()
                        candidate = re.sub(r'^[:\s]*|[:\s]*$', '', candidate)
                        # Strip coordinate position tags (#y84/L or [y=84|L]) before using as text
                        candidate = re.sub(r'#y\d+/[LCR]\s*|\[y=\d+\|[LCR]\]\s*', '', candidate).strip()
                        if not candidate:
                            break
                        if any(kw in candidate.upper() for kw in
                               ["NOTIFY", "SHIPPER", "VESSEL", "PORT OF"]):
                            break
                        parts.append(candidate)
                    if parts:
                        _consignee_parts = parts
                    break
            if _consignee_parts:
                break

        # Try 2: search near NOTIFY PARTIES for actual consignee details
        # Often the real consignee appears right after NOTIFY PARTIES on the B/L
        if not _consignee_parts:
            notify_idx = -1
            for i, line in enumerate(bl_lines):
                if "NOTIFY" in line.upper() and any(kw in line.upper() for kw in ["PARTY", "PARTIE", "PARTIES"]):
                    notify_idx = i
                    break
            if notify_idx >= 0:
                # Search lines after NOTIFY PARTIES for company name
                for j in range(notify_idx + 1, min(notify_idx + 10, len(bl_lines))):
                    candidate = bl_lines[j].strip()
                    candidate = re.sub(r'^[:\s]*|[:\s]*$', '', candidate)
                    candidate = re.sub(r'#y\d+/[LCR]\s*|\[y=\d+\|[LCR]\]\s*', '', candidate).strip()
                    if not candidate or len(candidate) < 3:
                        continue
                    # Skip parenthetical boilerplate clauses FIRST (before stop words check)
                    if candidate.startswith("(") or candidate.startswith(("(see", "(No ")):
                        continue
                    if any(kw in candidate.upper() for kw in
                           ["PORT OF", "VESSEL", "VOYAGE", "CONTAINER",
                            "FREIGHT", "SERVICE CONTRACT", "DISCHARGE"]):
                        break
                    # Accept company-like name lines
                    if re.search(r'\b[A-Z]{2,}\b', candidate) and candidate[0].isupper():
                        parts = [candidate]
                        for k in range(j + 1, min(j + 6, len(bl_lines))):
                            detail = bl_lines[k].strip()
                            detail = re.sub(r'^[:\s]*|[:\s]*$', '', detail)
                            detail = re.sub(r'#y\d+/[LCR]\s*|\[y=\d+\|[LCR]\]\s*', '', detail).strip()
                            if not detail or len(detail) < 3:
                                break
                            if any(kw in detail.upper() for kw in
                                   ["NOTIFY", "PORT OF", "VESSEL", "VOYAGE",
                                    "SERVICE CONTRACT", "DISCHARGE"]):
                                break
                            if re.search(r'[A-Za-z0-9]{2,}', detail):
                                parts.append(detail)
                        # Verify: not matching shipper
                        shipper_key = extracted.get("shipper", "")[:15].upper()
                        if not shipper_key or shipper_key not in candidate.upper():
                            _consignee_parts = parts
                        break

        # Apply consignee results: populate details always, name only if AI wrong/missing
        if _consignee_parts:
            fallback_name = _consignee_parts[0]
            fallback_details = "\n".join(_consignee_parts)
            shipper_upper = extracted.get("shipper", "").upper().strip()
            ai_consignee = (extracted.get("consignee") or "").strip()
            ai_upper = ai_consignee.upper()

            # Treat "TO ORDER" as placeholder, not a real consignee name
            is_to_order = any(kw in ai_upper for kw in ["TO ORDER", "TO THE ORDER"])

            if not ai_consignee or is_to_order:
                # AI returned nothing or "TO ORDER" → use fallback name + details
                extracted["consignee"] = fallback_name
                extracted["consignee_details"] = fallback_details
            elif ai_upper == shipper_upper:
                # AI returned shipper as consignee (wrong) → use fallback
                extracted["consignee"] = fallback_name
                extracted["consignee_details"] = fallback_details
            else:
                # AI returned a different name → keep AI's name, add details from regex
                # Only overwrite details if AI didn't return its own
                if not extracted.get("consignee_details"):
                    if fallback_name.upper().strip() != shipper_upper:
                        extracted["consignee_details"] = fallback_details
                # If AI name is longer/better than fallback first line, keep AI name
                if len(ai_consignee) < len(fallback_name):
                    # But only if fallback_name looks like a real company name
                    if fallback_name.upper().strip() != shipper_upper:
                        extracted["consignee"] = fallback_name

        # ═══════════════════════════════════════════════════════
        # Phase 2: Generic LLM-driven template filling
        # Works with ANY template format — no hardcoded mappings.
        # ═══════════════════════════════════════════════════════
        filled_docx_path = ""
        download_id = ""
        filled_letter = ""

        if "template" in saved_paths:
            template_path = saved_paths["template"]
            ext = os.path.splitext(template_path)[1].lower()

            # ── DOCX path: fast mapping first, then LLM for remaining gaps ──
            if ext == ".docx":
                # 2a. Fast path: hardcoded placeholder mapping (instant, works for
                #     [发货人抬头] / [发货人名称] / etc.)
                print("[auto-fill-telex] Phase 2: DOCX — hardcoded mapping first...")
                try:
                    filled_docx_path = await asyncio.to_thread(
                        _fill_telex_docx, template_path, extracted, carrier or ""
                    )
                    filled_letter = await asyncio.to_thread(
                        extract_text, filled_docx_path
                    )
                    download_id = str(uuid.uuid4())
                    _telex_download_store[download_id] = filled_docx_path
                    print(f"  Fast fill OK -> {os.path.basename(filled_docx_path)}")
                except Exception as e:
                    print(f"  Fast fill failed: {e}")
                    filled_docx_path = ""

                # 2b. Check if template still has unfilled placeholders / gaps
                _needs_llm = False
                if filled_docx_path:
                    _raw = filled_letter or ""
                    # Patterns that indicate something is still unfilled:
                    _unfilled_patterns = [
                        r'\[.*?\]',           # [发货人抬头] style
                        r'\.{4,}',            # ASCII dots ....
                        r'…{2,}',             # Unicode ellipsis ……
                        r'_{4,}',             # ____
                        r'\([^)]*name[^)]*\)',  # (shipper name) / (consignee's name)
                        r'\[insert[^\]]*\]',  # [insert ...]
                        r'shipper[^)]*name',  # "shipper's name" without parens
                        r'consignee[^)]*name',
                    ]
                    for pat in _unfilled_patterns:
                        if re.search(pat, _raw):
                            _needs_llm = True
                            break
                    print(f"  Unfilled gaps detected: {_needs_llm}")

                # 2c. LLM gap-fill (only if unfilled placeholders remain)
                if _needs_llm:
                    print("  Running LLM to fill remaining gaps...")
                    try:
                        # Find placeholder lines to send as context (much smaller)
                        _ph_lines = []
                        for line in _raw.split('\n'):
                            stripped = line.strip()
                            if not stripped or len(stripped) < 3:
                                continue
                            if re.search(r'\.{3,}|…{2,}|_{3,}|\([^)]*name[^)]*\)|\[insert[^\]]*\]', stripped):
                                _ph_lines.append(stripped)
                        _ph_text = "\n".join(_ph_lines[:20])
                        _sys = (
                            "You are a document-filling specialist. Replace the remaining "
                            "placeholders in this telex-release template. "
                            "Return ONLY a JSON array of {old, new} pairs."
                        )
                        _usr = (
                            f"## Unfilled placeholder lines\n{_ph_text}\n\n"
                            f"## Extracted B/L data\n{extracted_text}\n\n"
                            "Return JSON array: [{\"old\":\"EXACT text in the document\","
                            '"new":"replacement text"}]\n\n'
                            "Rules:\n"
                            "- Replace `.........`, `…………`, `____` with the actual data\n"
                            "- Replace `(shipper's name)` / `(consignee's name)` "
                            "with the actual names\n"
                            "- Replace `[insert ...]` style brackets with full details\n"
                            "- `[发货人抬头]` = shipper + address\n"
                            "- `[发货人名称]` = shipper only\n"
                            "- `[收货人详细信息]` = consignee_details\n"
                            "- `[收货人抬头]` = consignee only\n"
                            "- `old` must be the EXACT text including dots/ellipsis\n"
                            "- If no placeholder remains → return []\n"
                            "Return ONLY the JSON array."
                        )
                        p2_text = await _llm_call(
                            [{"role":"system","content":_sys},
                             {"role":"user","content":_usr}],
                            4096, 0.05
                        )
                        replacements = []
                        try:
                            replacements = json.loads(p2_text.strip())
                        except json.JSONDecodeError:
                            arr = re.search(r'\[.*\]', p2_text, re.DOTALL)
                            if arr:
                                try: replacements = json.loads(arr.group())
                                except: pass
                        if isinstance(replacements, dict):
                            replacements = replacements.get("replacements", [])
                        if isinstance(replacements, list) and replacements:
                            out_dir = os.path.join(upload_dir, "filled")
                            _gp = os.path.join(out_dir, f"telex_filled_{uuid.uuid4().hex}.docx")
                            await asyncio.to_thread(
                                _apply_global_replacements, filled_docx_path, _gp, replacements
                            )
                            filled_docx_path = _gp
                            filled_letter = await asyncio.to_thread(extract_text, _gp)
                            download_id = str(uuid.uuid4())
                            _telex_download_store[download_id] = _gp
                            print(f"  Gap-fill applied ({len(replacements)} replacements)")
                        else:
                            print(f"  LLM returned no usable replacements")
                    except Exception as e:
                        print(f"  LLM gap-fill failed: {e}")

            # ── DOC path: LibreOffice convert → DOCX pipeline (preserves format) ──
            elif ext == ".doc":
                print("[auto-fill-telex] Phase 2: .doc → LibreOffice convert + DOCX pipeline...")
                try:
                    # Convert .doc → .docx via LibreOffice headless (preserves
                    # tables, headers, fonts, layout — unlike text+LLM approach)
                    converted_docx_path = template_path + "_converted.docx"
                    soffice = "C:\\Program Files\\LibreOffice\\program\\soffice.exe"
                    convert_dir = os.path.dirname(converted_docx_path)
                    if not os.path.isdir(convert_dir):
                        convert_dir = upload_dir

                    async def _convert_doc_to_docx() -> str:
                        convert_in = os.path.join(convert_dir, f"convert_in_{uuid.uuid4().hex}.doc")
                        shutil.copy2(template_path, convert_in)
                        try:
                            proc = await asyncio.create_subprocess_exec(
                                soffice,
                                "--headless", "--convert-to", "docx",
                                convert_in,
                                "--outdir", convert_dir,
                                stdout=asyncio.subprocess.DEVNULL,
                                stderr=asyncio.subprocess.DEVNULL,
                            )
                            await asyncio.wait_for(proc.wait(), timeout=60)
                            converted = convert_in.replace(".doc", ".docx")
                            if os.path.exists(converted):
                                return converted
                            # soffice may lowercase the extension
                            converted_lower = convert_in.replace(".doc", ".docx")
                            if os.path.exists(converted_lower):
                                return converted_lower
                            # Search in convert_dir
                            for f in os.listdir(convert_dir):
                                if f.endswith(".docx") and f.startswith("convert_in_"):
                                    return os.path.join(convert_dir, f)
                            raise FileNotFoundError("LibreOffice did not produce output file")
                        finally:
                            for _p in [convert_in]:
                                try: os.remove(_p)
                                except: pass

                    converted_path = await _convert_doc_to_docx()
                    print(f"  Converted to DOCX: {os.path.basename(converted_path)} "
                          f"({os.path.getsize(converted_path)} bytes)")

                    # Now use the same DOCX pipeline: fast mapping + LLM gap-fill
                    try:
                        fill_out = await asyncio.to_thread(
                            _fill_telex_docx, converted_path, extracted, carrier or ""
                        )
                        filled_letter = await asyncio.to_thread(extract_text, fill_out)
                        filled_docx_path = fill_out
                        download_id = str(uuid.uuid4())
                        _telex_download_store[download_id] = fill_out
                        print(f"  Fast fill OK -> {os.path.basename(fill_out)}")
                    except Exception as e:
                        print(f"  Fast fill failed: {e}")

                    # Check for remaining gaps → regex gap-fill (reliable, no API)
                    if filled_docx_path:
                        _raw = filled_letter or ""
                        _needs_gap_fill = bool(re.search(
                            r'\.{3,}|…{2,}|_{3,}|\([^)]*name[^)]*\)|\[insert[^\]]*\]',
                            _raw
                        ))
                        if _needs_gap_fill:
                            print("  Regex gap-filling remaining placeholders...")
                            try:
                                _gp = fill_out.replace(".docx", f"_gapfill_{uuid.uuid4().hex}.docx")
                                await asyncio.to_thread(
                                    _regex_gap_fill, fill_out, _gp, extracted
                                )
                                filled_docx_path = _gp
                                filled_letter = await asyncio.to_thread(extract_text, _gp)
                                download_id = str(uuid.uuid4())
                                _telex_download_store[download_id] = _gp
                                print(f"  Regex gap-fill applied -> {os.path.basename(_gp)}")
                            except Exception as e:
                                print(f"  Regex gap-fill failed: {e}")
                        else:
                            print("  No remaining gaps detected")

                    # Clean up the temporary converted file
                    try: os.remove(converted_path)
                    except: pass

                except Exception as e:
                    print(f"[auto-fill-telex] .doc conversion+fill failed: {e}, falling back to text LLM fill...")
                    # Fallback: text+LLM (same as before, no format preservation)
                    try:
                        doc_text = await asyncio.to_thread(extract_text, template_path)
                        _sys = "You are a document-filling specialist. Fill the telex template using extracted data. Return ONLY filled text."
                        _usr = f"## Template\n{doc_text}\n\n## Data\n{extracted_text}\n\nCarrier: {carrier or 'Unknown'}\n\nFill ALL placeholders like [xxx], ...., ____, (shipper name). Output ONLY the filled text."
                        filled_letter = await _llm_call([{"role":"system","content":_sys},{"role":"user","content":_usr}], 4096)
                        out_path = template_path + "_filled.docx"
                        try:
                            from docx import Document as DocxDoc
                            d = DocxDoc()
                            for line in filled_letter.split('\n'): d.add_paragraph(line)
                            d.save(out_path)
                        except:
                            out_path = template_path + "_filled.txt"
                            with open(out_path,"w",encoding="utf-8") as f: f.write(filled_letter)
                        filled_docx_path = out_path
                        download_id = str(uuid.uuid4())
                        _telex_download_store[download_id] = out_path
                    except Exception as e3:
                        print(f"  Text fallback also failed: {e3}")

        else:
            # No template — use default text template with LLM fill
            print("[auto-fill-telex] Phase 2: no template — LLM fill default telex...")
            try:
                phase2_system = (
                    "You are a document-filling specialist. Fill in the telex "
                    "release template using the extracted B/L data."
                )
                phase2_user = (
                    f"## Template Text\n{template_text}\n\n"
                    f"## Extracted B/L Data (JSON)\n{extracted_text}\n\n"
                    f"## Carrier\n{carrier or 'Unknown'}\n\n"
                    "Fill the template. Replace _______________ with data."
                )
                p2_messages = [
                    {"role": "system", "content": phase2_system},
                    {"role": "user", "content": phase2_user},
                ]
                filled_letter = await _llm_call(p2_messages, max_tokens=4096)
            except Exception as e:
                print(f"[auto-fill-telex] Default template fill failed: {e}")
                filled_letter = _generate_telex_letter(carrier or "CARRIER", extracted)

        # If no DOCX was produced and we have text, clean it up
        if not filled_docx_path and filled_letter:
            filled_letter = re.sub(r'\[MISSING:[^\]]*\]', '_______________', filled_letter)
            filled_letter = re.sub(r"^={3,}.*?={3,}\s*", "", filled_letter, flags=re.MULTILINE).strip()
            _known_headers = [
                "Combined Letter of Indemnity", "Telex Release", "电放保函",
                "Letter of Indemnity", "TLX", "TELEX RELEASE",
                "Attach shipper Title Paper",
            ]
            for h in _known_headers:
                idx = filled_letter.rfind(h)
                if idx >= 0 and idx < len(filled_letter) - 50:
                    filled_letter = filled_letter[idx:].strip()
                    break

        return {
            "success": True,
            "filled_letter": filled_letter,
            "extracted": extracted,
            "download_id": download_id,
            "model_used": model_primary,
        }

    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        print(f"[auto-fill-telex] ERROR: {error_detail}")
        log_path = os.path.join(settings.UPLOAD_DIR, "auto_fill_telex_errors.log")
        try:
            with open(log_path, "a", encoding="utf-8") as logf:
                logf.write(f"\n[ERROR at {__import__('datetime').datetime.now()}]\n{error_detail}\n")
        except Exception:
            pass
        return {
            "success": False,
            "error": str(e),
            "filled_letter": "",
            "extracted": {},
            "model_used": "",
        }

    finally:
        for path in saved_paths.values():
            try:
                if os.path.exists(path):
                    os.remove(path)
            except Exception:
                pass


@router.get("/letter/download/{download_id}")
async def download_filled_telex(download_id: str):
    """下载已填写的电放保函 DOCX 文件。"""
    file_path = _telex_download_store.get(download_id)
    if not file_path or not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在或已过期")
    filename = f"电放保函_{uuid.uuid4().hex[:8]}.docx"
    from fastapi.responses import FileResponse
    return FileResponse(
        file_path,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        filename=filename,
    )
