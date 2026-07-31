"""
OAIW 文档解析模块 — 报价单/保函/箱单发票/舱单

支持格式: PDF, DOCX, XLSX, 图片
"""
from __future__ import annotations

import os
from typing import Optional


def extract_text(file_path: str) -> str:
    """根据文件扩展名提取文本内容。"""
    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".docx":
        return _extract_docx(file_path)
    elif ext == ".pdf":
        return _extract_pdf(file_path)
    elif ext == ".xlsx":
        return _extract_xlsx(file_path)
    elif ext in (".png", ".jpg", ".jpeg", ".bmp", ".webp"):
        return _extract_image(file_path)
    elif ext == ".doc":
        return _extract_doc(file_path)
    elif ext == ".txt":
        with open(file_path, "r", encoding="utf-8", errors="replace") as f:
            return f.read()
    else:
        return f"不支持的文件格式: {ext}"


def _extract_docx(file_path: str) -> str:
    """提取 DOCX 文本。"""
    import docx
    doc = docx.Document(file_path)
    lines = []
    for p in doc.paragraphs:
        if p.text.strip():
            lines.append(p.text.strip())
    for i, table in enumerate(doc.tables):
        lines.append(f"\n--- 表格 {i+1} ---")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            lines.append(" | ".join(cells))
    return "\n".join(lines)


def _extract_doc(file_path: str) -> str:
    """提取旧版 .doc 文本 (使用 antiword CLI 或 olefile 回退)。"""
    import shutil, subprocess

    # 1) antiword — 提取效果最好
    # 在 Windows 下 Git for Windows 会安装 antiword 到 mingw64/bin
    # 注意: antiword 对含特殊字符的文件名可能失败，先复制到临时路径
    import tempfile
    antiword_paths = [
        shutil.which("antiword"),
        "C:\\Program Files\\Git\\mingw64\\bin\\antiword.exe",
        "C:\\Program Files (x86)\\Git\\mingw64\\bin\\antiword.exe",
    ]
    for antiword in filter(None, antiword_paths):
        if os.path.exists(antiword):
            try:
                # 复制到临时文件以避免 antiword 对特殊字符文件名不兼容
                tmp_fd, tmp_path = tempfile.mkstemp(suffix=".doc")
                os.close(tmp_fd)
                shutil.copy2(file_path, tmp_path)
                res = subprocess.run(
                    [antiword, tmp_path],
                    capture_output=True, timeout=15,
                    env={**os.environ, "HOME": os.environ.get("USERPROFILE", "C:\\")},
                )
                os.unlink(tmp_path)
                if res.returncode == 0 and res.stdout and len(res.stdout) > 100:
                    # antiword 输出混合二进制，尝试提取纯文本行
                    import re
                    raw = res.stdout
                    text = raw.decode("utf-8", errors="replace")
                    lines = [l.strip() for l in text.replace('\r\n', '\n').split('\n') if l.strip()]
                    # 保留包含英文字母/数字的文本行（过滤纯二进制行）
                    meaningful = [
                        l for l in lines
                        if len(l) > 3 and (re.search(r'[A-Za-z]{2,}', l) or re.search(r'[一-鿿]', l))
                    ]
                    if meaningful:
                        return '\n'.join(meaningful)
            except Exception:
                pass

    # 2) catdoc 回退
    for catdoc in [shutil.which("catdoc")]:
        if catdoc:
            try:
                res = subprocess.run(
                    [catdoc, file_path],
                    capture_output=True, timeout=15,
                )
                if res.returncode == 0 and res.stdout:
                    text = res.stdout.decode("utf-8", errors="replace")
                    if text.strip():
                        return text.strip()
            except Exception:
                pass

    # 3) olefile 极简回退 — 读取 WordDocument / Text 流
    try:
        import olefile
        ole = olefile.OleFileIO(file_path)
        for stream_name in ("WordDocument", "1Table", "0Table"):
            if ole.exists(stream_name):
                data = ole.openstream(stream_name).read()
                # 提取可打印 ASCII / 中文
                text = "".join(chr(b) for b in data if 32 <= b < 127 or b >= 160)
                if text.strip():
                    ole.close()
                    return text.strip()
        ole.close()
    except Exception:
        pass

    return f"[无法提取 .doc 文件内容: {os.path.basename(file_path)}]"


def _extract_pdf(file_path: str) -> str:
    """提取 PDF 文本。扫描件自动降级为 OCR。"""
    import fitz
    doc = fitz.open(file_path)
    lines = []
    has_text = False
    for page in doc:
        text = page.get_text().strip()
        if text:
            has_text = True
            lines.append(f"--- 第 {page.number + 1} 页 ---")
            lines.append(text)

    if has_text:
        return "\n".join(lines)

    # 扫描件 OCR 回退
    return _ocr_pdf_pages(doc)


def _ocr_pdf_pages(doc) -> str:
    """将 PDF 每页渲染为图片后 OCR。

    使用 PaddleOCR-json 引擎（通过 Umi-OCR 的 PPOCR_pipe）
    - 中文识别率远优于 easyocr CPU 模式
    - 支持中英混合识别
    """
    import os, sys, tempfile, shutil
    from pathlib import Path

    # PaddleOCR-json 引擎路径
    _PADDLE_EXE = os.path.join(
        "D:/Umi-OCR/Umi-OCR_Paddle_v2.1.5/UmiOCR-data/plugins",
        "win7_x64_PaddleOCR-json/PaddleOCR-json.exe",
    )
    _PPOCR_API = os.path.join(
        "D:/Umi-OCR/Umi-OCR_Paddle_v2.1.5/UmiOCR-data/plugins",
        "win7_x64_PaddleOCR-json",
    )

    if not os.path.exists(_PADDLE_EXE):
        # Fallback: try easyocr
        return _ocr_fallback_easyocr(doc)

    # Add PPOCR_api.py to path
    sys.path.insert(0, _PPOCR_API)
    from PPOCR_api import PPOCR_pipe

    ocr = None
    tmpdir = None
    try:
        ocr = PPOCR_pipe(_PADDLE_EXE)
        tmpdir = tempfile.mkdtemp()
        all_text = []

        for page_num in range(doc.page_count):
            page = doc[page_num]
            pix = page.get_pixmap(dpi=600)
            img_path = os.path.join(tmpdir, f"p{page_num}.png")
            pix.save(img_path)

            res = ocr.run(img_path)
            if res.get("code") == 100:
                texts = [b.get("text", "") for b in (res.get("data") or [])]
                if texts:
                    all_text.append(f"--- 第 {page_num + 1} 页 ---")
                    all_text.extend(texts)
            else:
                all_text.append(f"--- 第 {page_num + 1} 页 ---")
                all_text.append(f"[OCR warn: {res.get('data', 'unknown')}]")

        return "\n".join(all_text) if all_text else "[OCR 未识别到文字]"

    except Exception as e:
        return f"[PaddleOCR 引擎错误: {e}]"
    finally:
        if ocr:
            try:
                ocr.exit()
            except Exception:
                pass
        if tmpdir and os.path.exists(tmpdir):
            try:
                shutil.rmtree(tmpdir)
            except Exception:
                pass


def _ocr_fallback_easyocr(doc) -> str:
    """回退到 easyocr（PaddleOCR-json 不可用时）。"""
    try:
        import easyocr

        reader = easyocr.Reader(["ch_sim", "en"], gpu=False)
    except ImportError:
        return "[PDF 为扫描件且 OCR 引擎未安装]"

    all_text = []
    for page_num in range(doc.page_count):
        page = doc[page_num]
        pix = page.get_pixmap(dpi=600)
        img_bytes = pix.tobytes("png")
        results = reader.readtext(img_bytes, paragraph=True)
        page_lines = [r[-1] for r in results if r and r[-1]]
        if page_lines:
            all_text.append(f"--- 第 {page_num + 1} 页 ---")
            all_text.extend(page_lines)

    return "\n".join(all_text) if all_text else "[OCR 未识别到文字]"


def _extract_xlsx(file_path: str) -> str:
    """提取 XLSX 内容。"""
    import openpyxl
    wb = openpyxl.load_workbook(file_path, data_only=True)
    lines = []
    for sheet_name in wb.sheetnames:
        ws = wb[sheet_name]
        lines.append(f"\n--- 工作表: {sheet_name} ---")
        for row in ws.iter_rows(values_only=True):
            cells = [str(c) if c is not None else "" for c in row]
            line = "\t".join(cells)
            if line.strip():
                lines.append(line)
    return "\n".join(lines)


def _extract_image(file_path: str) -> str:
    """图片 OCR 提取文字。"""
    try:
        import paddleocr
        ocr = paddleocr.PaddleOCR(use_angle_cls=True, lang="ch", use_gpu=False, show_log=False)
        result = ocr.ocr(file_path, cls=True)
        lines = []
        if result and result[0]:
            for line in result[0]:
                text = line[1][0] if line[1] else ""
                if text.strip():
                    lines.append(text.strip())
        return "\n".join(lines) if lines else "[OCR 未识别到文字]"
    except ImportError:
        # PaddleOCR 未安装，回退到 easyocr
        try:
            import easyocr
            reader = easyocr.Reader(["ch_sim", "en"], gpu=False)
            results = reader.readtext(file_path, paragraph=True)
            return "\n".join(r[-1] for r in results if r and r[-1])
        except ImportError:
            return "[OCR 引擎未安装，请安装 paddleocr 或 easyocr]"
    except Exception as e:
        return f"[OCR 识别失败: {str(e)}]"
