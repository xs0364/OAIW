import sys, io, subprocess, os, uuid, shutil

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

src = r"D:\OAIW\dianfangbaohan\2\template_case2.doc"
conv = r"D:\OAIW\temp_lo2"
os.makedirs(conv, exist_ok=True)
tmp = os.path.join(conv, f"check_{uuid.uuid4().hex}.doc")
shutil.copy2(src, tmp)
subprocess.run(["C:\\Program Files\\LibreOffice\\program\\soffice.exe",
    "--headless","--convert-to","docx",tmp,"--outdir",conv], capture_output=True, timeout=60)
conv_file = tmp.replace(".doc", ".docx")
if os.path.exists(conv_file):
    from docx import Document
    import xml.etree.ElementTree as ET
    doc = Document(conv_file)

    print("=== PARAGRAPHS with labels ===")
    for i, p in enumerate(doc.paragraphs):
        t = p.text
        # Check for any circled numbers or (N) patterns
        has_label = any(ch in t for ch in '①②③④⑤⑥⑦⑧⑨⑩⑪⑫⑬⑭⑮⑯⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃')
        has_paren_num = bool(__import__('re').search(r'\(\d+\)', t))
        if has_label or has_paren_num:
            print(f"  [{i:2d}] {repr(t[:300])}")

    print("\n=== HEADERS ===")
    for si, s in enumerate(doc.sections):
        for hdr in [s.header]:
            if hdr:
                for p in hdr.paragraphs:
                    if p.text.strip():
                        print(f"  HDR(S{si}): {repr(p.text[:200])}")

    # Cleanup
    for f in os.listdir(conv):
        os.remove(os.path.join(conv, f))
    os.rmdir(conv)
